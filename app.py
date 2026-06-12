from __future__ import annotations

import io
import json
import re
import time
import zipfile
from dataclasses import replace
from datetime import datetime
from pathlib import Path
from typing import Callable

import cv2
import pandas as pd
import streamlit as st
from PIL import Image

from src.detector import Detection, ImageAnalysis, analyze_image, annotate_detections
from src.reporting import build_summary_table, load_sample_metadata, sequence_rows_from_uploads
from src.vigor import compute_vigor_metrics, score_vigor


ROOT = Path(__file__).resolve().parent
MODEL_PATH = ROOT / "models" / "seed_detector.pt"
RESULTS_DIR = ROOT / "results"
TRAINING_RESULTS = ROOT / "runs" / "seed_yolov8s_zm_full" / "results.csv"
SAMPLE_DIR = ROOT / "data" / "sample_sequence"
SAMPLE_META = SAMPLE_DIR / "metadata.csv"
CAPTURE_DIR = ROOT / "captures"


st.set_page_config(page_title="SeedGerm-Vigor", page_icon="🌱", layout="wide")


def inject_theme() -> None:
    st.markdown(
        """
        <style>
        :root {
            --bg: #f7f8f4;
            --panel: #ffffff;
            --border: #d9e2d8;
            --text: #1f2933;
            --muted: #61706b;
            --accent: #16846f;
            --accent-dark: #0e5f51;
        }
        .stApp { background: var(--bg); color: var(--text); }
        .main .block-container { max-width: 1280px; padding-top: 1.4rem; padding-bottom: 2.5rem; }
        [data-testid="stSidebar"] { background: #eef3ec; border-right: 1px solid var(--border); }
        [data-testid="stSidebar"] h3 { color: var(--accent-dark); font-size: 0.95rem; }
        .app-header {
            display: grid; grid-template-columns: minmax(0, 1fr) 360px;
            gap: 1rem; align-items: stretch; margin-bottom: 1.1rem;
        }
        .title-panel h1 {
            color: var(--text); font-size: clamp(1.75rem, 2.5vw, 2.45rem);
            line-height: 1.15; margin: 0; letter-spacing: 0;
        }
        .title-panel p {
            color: var(--muted); font-size: 0.98rem; line-height: 1.55;
            margin: 0.55rem 0 0; max-width: 780px;
        }
        .eyebrow {
            color: var(--accent); font-size: 0.78rem; font-weight: 700;
            letter-spacing: 0; margin-bottom: 0.25rem; text-transform: uppercase;
        }
        .status-panel {
            background: var(--panel); border: 1px solid var(--border); border-radius: 8px;
            padding: 0.95rem 1rem; box-shadow: 0 8px 22px rgba(31, 41, 51, 0.05);
        }
        .status-row {
            display: flex; align-items: center; justify-content: space-between; gap: 0.75rem;
            border-bottom: 1px solid #edf1eb; padding: 0.45rem 0;
        }
        .status-row:first-child { padding-top: 0; }
        .status-row:last-child { border-bottom: 0; padding-bottom: 0; }
        .status-label { color: var(--muted); font-size: 0.8rem; }
        .status-value { color: var(--text); font-size: 0.86rem; font-weight: 700; text-align: right; }
        .mode-chip {
            display: inline-flex; align-items: center; border-radius: 999px;
            border: 1px solid #b9d6cd; background: #e8f5f0; color: var(--accent-dark);
            font-size: 0.78rem; font-weight: 700; padding: 0.18rem 0.55rem;
        }
        h2 { font-size: 1.22rem !important; margin-top: 1.2rem !important;
             padding-bottom: 0.35rem; border-bottom: 1px solid var(--border); }
        h3 { font-size: 1.02rem !important; }
        [data-testid="stMetric"] {
            background: var(--panel); border: 1px solid var(--border); border-radius: 8px;
            padding: 0.85rem 0.95rem; box-shadow: 0 6px 18px rgba(31, 41, 51, 0.04);
        }
        [data-testid="stMetricLabel"] p { color: var(--muted); font-size: 0.78rem; }
        [data-testid="stMetricValue"] { color: var(--text); font-size: 1.45rem; }
        .stDataFrame { border: 1px solid var(--border); border-radius: 8px; overflow: hidden; }
        div[data-testid="stFileUploader"] {
            background: var(--panel); border: 1px dashed #aac4ba; border-radius: 8px; padding: 0.3rem 0.8rem;
        }
        div.stButton > button, div.stDownloadButton > button {
            border-radius: 8px; border: 1px solid var(--accent); background: var(--accent);
            color: #ffffff; font-weight: 700; min-height: 2.55rem;
        }
        div.stButton > button:hover, div.stDownloadButton > button:hover {
            border-color: var(--accent-dark); background: var(--accent-dark); color: #ffffff;
        }
        @media (max-width: 900px) { .app-header { grid-template-columns: 1fr; } }
        </style>
        """,
        unsafe_allow_html=True,
    )


def render_app_header(model_status: str, analyzer_name: str, mode: str) -> None:
    st.markdown(
        f"""
        <div class="app-header">
          <div class="title-panel">
            <div class="eyebrow">SeedGerm-Vigor</div>
            <h1>种子活力视觉表型监测与自动评价系统</h1>
            <p>面向培养皿图像和农业发芽实验，提供低成本图像采集、种子检测、时序趋势分析、人工校正闭环和实验报告输出。</p>
          </div>
          <div class="status-panel">
            <div class="status-row"><span class="status-label">模型状态</span><span class="status-value">{model_status}</span></div>
            <div class="status-row"><span class="status-label">推理模式</span><span class="status-value">{analyzer_name}</span></div>
            <div class="status-row"><span class="status-label">当前任务</span><span class="mode-chip">{mode}</span></div>
          </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


inject_theme()


@st.cache_resource(show_spinner="正在加载 YOLO 检测模型...")
def load_yolo_detector(weights_path: str, weights_mtime: float):
    from src.yolo_detector import YOLOSeedDetector

    return YOLOSeedDetector(weights_path)


@st.cache_data
def read_csv_if_exists(path: str) -> pd.DataFrame | None:
    csv_path = Path(path)
    if not csv_path.exists():
        return None
    return pd.read_csv(csv_path)


def get_analyzer(detector_choice: str, conf: float, iou: float) -> tuple[Callable[[Image.Image], ImageAnalysis], str]:
    if detector_choice.startswith("YOLO") and MODEL_PATH.exists():
        try:
            detector = load_yolo_detector(str(MODEL_PATH), MODEL_PATH.stat().st_mtime)
            detector.conf = conf
            detector.iou = iou
            return detector.analyze, detector.mode_name
        except Exception as exc:  # pragma: no cover - Streamlit fallback path
            st.sidebar.error("YOLO 模型加载失败，已切换到 OpenCV 基线。")
            st.sidebar.caption(str(exc))
    return analyze_image, "OpenCV rule-based baseline"


def postprocess_analysis(
    analysis: ImageAnalysis,
    conservative_germination: bool,
    germination_conf_threshold: float,
    germination_aspect_threshold: float,
) -> ImageAnalysis:
    if not conservative_germination:
        return analysis

    detections = []
    for detection in analysis.detections:
        if detection.label != "germinated":
            detections.append(detection)
            continue
        high_confidence = detection.confidence >= germination_conf_threshold
        elongated_shape = detection.aspect_ratio >= germination_aspect_threshold
        detections.append(detection if high_confidence or elongated_shape else replace(detection, label="non_germinated"))
    return ImageAnalysis(detections=detections)


def show_single_image_analysis(
    image: Image.Image,
    caption: str,
    analyzer: Callable[[Image.Image], ImageAnalysis],
) -> ImageAnalysis:
    result = analyzer(image)
    annotated = annotate_detections(image, result.detections)

    left, right = st.columns([1.25, 1])
    with left:
        st.image(annotated, caption=caption, use_column_width=True)
    with right:
        c1, c2 = st.columns(2)
        c1.metric("检测种子数", result.total)
        c2.metric("发芽率", f"{result.germination_rate:.1f}%")
        c3, c4 = st.columns(2)
        c3.metric("已发芽", result.germinated)
        c4.metric("未发芽", result.non_germinated)
        st.dataframe(result.to_dataframe(), hide_index=True, use_container_width=True)
    return result


def render_vigor_summary(summary: pd.DataFrame) -> None:
    metrics = compute_vigor_metrics(summary)
    vigor_score = score_vigor(metrics)
    top = st.columns(4)
    top[0].metric("最终发芽率", f"{metrics.final_germination_rate:.1f}%")
    top[1].metric("T50", "未达到" if metrics.t50_h is None else f"{metrics.t50_h:.1f} h")
    top[2].metric("平均发芽时间", f"{metrics.mean_germination_time_h:.1f} h")
    top[3].metric("活力评分", f"{vigor_score:.1f} / 100")

    if not summary.empty:
        st.line_chart(summary.set_index("time_h")[["germination_rate"]])
        st.dataframe(summary, hide_index=True, use_container_width=True)


def pretty_frame(frame: pd.DataFrame, columns: dict[str, str]) -> pd.DataFrame:
    visible = [column for column in columns if column in frame.columns]
    return frame[visible].rename(columns=columns)


def parse_batch_upload_name(name: str) -> tuple[str, float | None]:
    normalized = name.replace("\\", "/")
    parts = [part for part in normalized.split("/") if part]
    stem = Path(parts[-1]).stem if parts else Path(name).stem
    time_match = re.search(r"(\d+(?:\.\d+)?)\s*(?:h|hr|hour|hours|小时)", stem, flags=re.IGNORECASE)
    time_h = float(time_match.group(1)) if time_match else None

    if len(parts) >= 2:
        batch = parts[-2]
    elif time_match:
        batch = stem[: time_match.start()]
    else:
        batch = re.sub(r"[_-]?\d+.*$", "", stem)
    return re.sub(r"[_\-\s]+$", "", batch).strip() or "Batch 1", time_h


def analyze_batch_uploads(
    uploads,
    interval_hours: float,
    analyzer: Callable[[Image.Image], ImageAnalysis],
) -> tuple[pd.DataFrame, pd.DataFrame]:
    file_rows = []
    for upload in uploads:
        batch, time_h = parse_batch_upload_name(upload.name)
        file_rows.append(
            {"batch": batch, "filename": upload.name, "parsed_time_h": time_h, "image": Image.open(upload).convert("RGB")}
        )

    detail_records = []
    for batch, rows in pd.DataFrame(file_rows).groupby("batch", sort=True):
        batch_rows = rows.to_dict("records")
        batch_rows.sort(key=lambda row: (row["parsed_time_h"] is None, row["parsed_time_h"] or 0.0, row["filename"]))
        for index, row in enumerate(batch_rows):
            time_h = row["parsed_time_h"]
            if time_h is None:
                time_h = round(index * float(interval_hours), 2)
            result = analyzer(row["image"])
            detail_records.append(
                {
                    "batch": batch,
                    "time_h": float(time_h),
                    "filename": row["filename"],
                    "total": result.total,
                    "germinated": result.germinated,
                    "non_germinated": result.non_germinated,
                    "germination_rate": round(result.germination_rate, 2),
                }
            )

    details = pd.DataFrame(detail_records).sort_values(["batch", "time_h", "filename"]).reset_index(drop=True)
    comparison_records = []
    for batch, summary in details.groupby("batch", sort=True):
        metrics = compute_vigor_metrics(summary)
        final = summary.sort_values("time_h").iloc[-1]
        comparison_records.append(
            {
                "batch": batch,
                "time_points": len(summary),
                "final_total": int(final["total"]),
                "final_germinated": int(final["germinated"]),
                "final_germination_rate": round(metrics.final_germination_rate, 2),
                "t50_h": None if metrics.t50_h is None else round(metrics.t50_h, 2),
                "mean_germination_time_h": round(metrics.mean_germination_time_h, 2),
                "germination_speed_index": round(metrics.germination_speed_index, 3),
                "uniformity_score": round(metrics.uniformity_score, 2),
                "vigor_score": round(score_vigor(metrics), 2),
            }
        )

    comparison = pd.DataFrame(comparison_records).sort_values(["vigor_score", "final_germination_rate"], ascending=False)
    return comparison.reset_index(drop=True), details


def build_markdown_report(experiment_name: str, comparison: pd.DataFrame, details: pd.DataFrame) -> str:
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    lines = [
        f"# {experiment_name}",
        "",
        f"- 生成时间：{generated_at}",
        f"- 批次数量：{comparison['batch'].nunique() if not comparison.empty else 0}",
        f"- 图像数量：{len(details)}",
        f"- 检测模型：{'YOLO 训练检测器' if MODEL_PATH.exists() else 'OpenCV 规则基线'}",
        "",
    ]
    if not comparison.empty:
        best = comparison.iloc[0]
        lines.extend(
            [
                "## 自动结论",
                "",
                f"综合活力评分最高的批次为 **{best['batch']}**，评分为 **{best['vigor_score']} / 100**。",
                f"其最终发芽率为 **{best['final_germination_rate']}%**。",
                "",
                "## 批次对比表",
                "",
                comparison.to_markdown(index=False),
                "",
            ]
        )
    lines.extend(["## 时间点明细", "", details.to_markdown(index=False) if not details.empty else "无数据", ""])
    return "\n".join(lines)


def build_docx_report(experiment_name: str, comparison: pd.DataFrame, details: pd.DataFrame) -> bytes | None:
    try:
        from docx import Document
    except Exception:
        return None

    document = Document()
    document.add_heading(experiment_name, level=1)
    document.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"图像数量：{len(details)}")
    document.add_paragraph(f"批次数量：{comparison['batch'].nunique() if not comparison.empty else 0}")
    document.add_paragraph(f"检测模型：{'YOLO 训练检测器' if MODEL_PATH.exists() else 'OpenCV 规则基线'}")
    if not comparison.empty:
        best = comparison.iloc[0]
        document.add_heading("自动结论", level=2)
        document.add_paragraph(
            f"综合活力评分最高的批次为 {best['batch']}，评分为 {best['vigor_score']} / 100，"
            f"最终发芽率为 {best['final_germination_rate']}%。"
        )

    for title, frame in [("批次对比表", comparison), ("时间点明细", details)]:
        document.add_heading(title, level=2)
        if frame.empty:
            document.add_paragraph("无数据")
            continue
        table = document.add_table(rows=1, cols=len(frame.columns))
        table.style = "Table Grid"
        for cell, column in zip(table.rows[0].cells, frame.columns):
            cell.text = str(column)
        for _, row in frame.iterrows():
            cells = table.add_row().cells
            for cell, value in zip(cells, row.tolist()):
                cell.text = "" if pd.isna(value) else str(value)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_report_package(experiment_name: str, comparison: pd.DataFrame, details: pd.DataFrame) -> bytes:
    markdown_report = build_markdown_report(experiment_name, comparison, details)
    docx_report = build_docx_report(experiment_name, comparison, details)
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("batch_comparison.csv", comparison.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("timepoint_details.csv", details.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("experiment_report.md", markdown_report.encode("utf-8"))
        if docx_report is not None:
            archive.writestr("experiment_report.docx", docx_report)
    return buffer.getvalue()


def capture_usb_frame(camera_index: int) -> Image.Image:
    cap = cv2.VideoCapture(int(camera_index), cv2.CAP_DSHOW)
    if not cap.isOpened():
        cap = cv2.VideoCapture(int(camera_index))
    ok, frame = cap.read()
    cap.release()
    if not ok or frame is None:
        raise RuntimeError(f"无法读取摄像头 {camera_index}，请检查设备编号或权限。")
    rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    return Image.fromarray(rgb)


def append_capture_record(
    image: Image.Image,
    analyzer: Callable[[Image.Image], ImageAnalysis],
    session_dir: Path,
    frame_index: int,
    start_time: float,
) -> dict:
    analysis = analyzer(image)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    image_name = f"frame_{frame_index:03d}_{timestamp}.jpg"
    image_path = session_dir / image_name
    image.save(image_path, quality=92)
    return {
        "frame": frame_index,
        "time_h": round((time.time() - start_time) / 3600.0, 4),
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "image_path": str(image_path),
        "total": analysis.total,
        "germinated": analysis.germinated,
        "non_germinated": analysis.non_germinated,
        "germination_rate": round(analysis.germination_rate, 2),
    }


def render_timed_camera_page(analyzer: Callable[[Image.Image], ImageAnalysis]) -> None:
    st.header("摄像头定时采集与自动趋势更新")
    st.caption("连接普通 USB 摄像头后，系统可按设定间隔自动采集培养皿图像，并实时更新发芽率曲线和活力指标。")

    if "timed_records" not in st.session_state:
        st.session_state.timed_records = []

    c1, c2, c3, c4 = st.columns(4)
    session_name = c1.text_input("实验编号", value=f"session_{datetime.now().strftime('%Y%m%d')}")
    camera_index = c2.number_input("摄像头编号", min_value=0, max_value=10, value=0, step=1)
    interval_seconds = c3.number_input("采集间隔（秒）", min_value=2, max_value=3600, value=10, step=1)
    frame_count = c4.number_input("采集张数", min_value=1, max_value=200, value=6, step=1)

    b1, b2, b3 = st.columns([1, 1, 3])
    start = b1.button("开始定时采集")
    manual = b2.button("采集单帧")
    if b3.button("清空当前监测记录"):
        st.session_state.timed_records = []
        st.rerun()

    image_box = st.empty()
    chart_box = st.empty()
    table_box = st.empty()

    if start or manual:
        session_dir = CAPTURE_DIR / f"{session_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        session_dir.mkdir(parents=True, exist_ok=True)
        start_time = time.time()
        n_frames = 1 if manual else int(frame_count)
        progress = st.progress(0.0)

        for frame_index in range(1, n_frames + 1):
            try:
                image = capture_usb_frame(int(camera_index))
            except Exception as exc:
                st.error(str(exc))
                break
            record = append_capture_record(image, analyzer, session_dir, frame_index, start_time)
            st.session_state.timed_records.append(record)
            analysis = ImageAnalysis(
                [
                    Detection(0, 0, 1, 1, 1.0, 1.0, "germinated", 1.0)
                    for _ in range(int(record["germinated"]))
                ]
                + [
                    Detection(0, 0, 1, 1, 1.0, 1.0, "non_germinated", 1.0)
                    for _ in range(int(record["non_germinated"]))
                ]
            )
            image_box.image(image, caption=f"最新采集：{record['timestamp']} | 发芽率 {analysis.germination_rate:.1f}%", use_column_width=True)

            df = pd.DataFrame(st.session_state.timed_records)
            chart_box.line_chart(df.set_index("time_h")[["germination_rate"]])
            table_box.dataframe(df, hide_index=True, use_container_width=True)
            progress.progress(frame_index / n_frames)
            if start and frame_index < n_frames:
                time.sleep(float(interval_seconds))

    if st.session_state.timed_records:
        records = pd.DataFrame(st.session_state.timed_records)
        st.subheader("实时趋势与活力指标")
        render_vigor_summary(records)
        csv = records.to_csv(index=False).encode("utf-8-sig")
        st.download_button("下载定时采集记录 CSV", data=csv, file_name=f"{session_name}_camera_records.csv", mime="text/csv")
    else:
        st.info("尚无定时采集记录。固定摄像头视野后，点击“开始定时采集”即可。")


def correction_frame_from_analysis(analysis: ImageAnalysis) -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "include": True,
                "x": int(det.x),
                "y": int(det.y),
                "w": int(det.w),
                "h": int(det.h),
                "label": det.label,
                "confidence": round(float(det.confidence), 3),
            }
            for det in analysis.detections
        ]
    )


def analysis_from_corrections(frame: pd.DataFrame) -> ImageAnalysis:
    detections = []
    for _, row in frame.iterrows():
        if not bool(row.get("include", True)):
            continue
        w = max(1, int(row["w"]))
        h = max(1, int(row["h"]))
        detections.append(
            Detection(
                x=max(0, int(row["x"])),
                y=max(0, int(row["y"])),
                w=w,
                h=h,
                area=float(w * h),
                aspect_ratio=float(max(w / max(1, h), h / max(1, w))),
                label=str(row["label"]),
                confidence=float(row.get("confidence", 1.0)),
            )
        )
    return ImageAnalysis(detections=detections)


def yolo_label_lines(frame: pd.DataFrame, width: int, height: int) -> list[str]:
    lines = []
    for _, row in frame.iterrows():
        if not bool(row.get("include", True)):
            continue
        class_id = 1 if str(row["label"]) == "germinated" else 0
        x = max(0.0, float(row["x"]))
        y = max(0.0, float(row["y"]))
        w = max(1.0, float(row["w"]))
        h = max(1.0, float(row["h"]))
        xc = min(1.0, max(0.0, (x + w / 2.0) / width))
        yc = min(1.0, max(0.0, (y + h / 2.0) / height))
        bw = min(1.0, max(0.0, w / width))
        bh = min(1.0, max(0.0, h / height))
        lines.append(f"{class_id} {xc:.6f} {yc:.6f} {bw:.6f} {bh:.6f}")
    return lines


def build_correction_package(image: Image.Image, image_name: str, edited: pd.DataFrame) -> bytes:
    image_stem = Path(image_name).stem or "corrected_image"
    corrected = edited.loc[edited["include"].astype(bool)].copy()
    metadata = {
        "image": image_name,
        "width": image.width,
        "height": image.height,
        "exported_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "classes": {"0": "non_germinated", "1": "germinated"},
        "count": int(len(corrected)),
    }

    image_buffer = io.BytesIO()
    image.save(image_buffer, format="JPEG", quality=92)
    label_text = "\n".join(yolo_label_lines(corrected, image.width, image.height)) + "\n"

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(f"images/{image_stem}.jpg", image_buffer.getvalue())
        archive.writestr(f"labels/{image_stem}.txt", label_text.encode("utf-8"))
        archive.writestr("corrections.csv", edited.to_csv(index=False).encode("utf-8-sig"))
        archive.writestr("metadata.json", json.dumps(metadata, ensure_ascii=False, indent=2).encode("utf-8"))
    return buffer.getvalue()


def render_correction_page(analyzer: Callable[[Image.Image], ImageAnalysis]) -> None:
    st.header("人工校正与再训练数据导出")
    st.caption("用户可修正误检、漏检和类别错误；系统会重新计算发芽率，并导出 YOLO 格式标签用于后续微调。")

    uploaded = st.file_uploader("上传需要校正的培养皿图像", type=["jpg", "jpeg", "png", "bmp", "webp"], key="correction_upload")
    if uploaded is None:
        st.info("上传图像后，先由模型生成初始检测框，再在表格中修改 include、类别和框位置。")
        return

    image = Image.open(uploaded).convert("RGB")
    analysis = analyzer(image)
    annotated = annotate_detections(image, analysis.detections)

    left, right = st.columns([1.15, 1])
    with left:
        st.image(annotated, caption="模型初始检测结果", use_column_width=True)
    with right:
        st.metric("模型检测种子数", analysis.total)
        st.metric("模型发芽率", f"{analysis.germination_rate:.1f}%")

    initial = correction_frame_from_analysis(analysis)
    if initial.empty:
        initial = pd.DataFrame(columns=["include", "x", "y", "w", "h", "label", "confidence"])

    edited = st.data_editor(
        initial,
        hide_index=True,
        use_container_width=True,
        num_rows="dynamic",
        column_config={
            "include": st.column_config.CheckboxColumn("保留", default=True),
            "label": st.column_config.SelectboxColumn("类别", options=["germinated", "non_germinated"]),
            "confidence": st.column_config.NumberColumn("置信度", min_value=0.0, max_value=1.0, step=0.01),
        },
    )

    corrected_analysis = analysis_from_corrections(edited)
    st.subheader("校正后统计")
    cols = st.columns(4)
    cols[0].metric("校正后总数", corrected_analysis.total)
    cols[1].metric("已发芽", corrected_analysis.germinated)
    cols[2].metric("未发芽", corrected_analysis.non_germinated)
    cols[3].metric("发芽率", f"{corrected_analysis.germination_rate:.1f}%")

    corrected_image = annotate_detections(image, corrected_analysis.detections)
    st.image(corrected_image, caption="校正后结果预览", use_column_width=True)

    package = build_correction_package(image, uploaded.name, edited)
    st.download_button(
        "导出再训练数据包（图像 + YOLO 标签 + 校正表）",
        data=package,
        file_name=f"{Path(uploaded.name).stem}_correction_dataset.zip",
        mime="application/zip",
    )


def render_batch_comparison_page(analyzer: Callable[[Image.Image], ImageAnalysis], interval_hours: float) -> None:
    st.header("多批次对比与实验报告导出")
    st.caption("文件名建议使用 `batchA_00h.jpg`、`batchA_06h.jpg`、`batchB_00h.jpg`。")

    experiment_name = st.text_input("实验名称", value="种子发芽活力对比实验")
    uploads = st.file_uploader(
        "上传多个批次的时序图像",
        type=["jpg", "jpeg", "png", "bmp", "webp"],
        accept_multiple_files=True,
    )
    if not uploads:
        st.info("上传后会自动生成批次对比表、发芽率曲线和可下载实验报告。")
        return

    with st.spinner("正在逐张图像检测并计算批次指标..."):
        comparison, details = analyze_batch_uploads(uploads, interval_hours, analyzer)

    st.subheader("批次综合对比")
    comparison_view = pretty_frame(
        comparison,
        {
            "batch": "批次",
            "time_points": "时间点数",
            "final_total": "末次总数",
            "final_germinated": "末次已发芽",
            "final_germination_rate": "最终发芽率(%)",
            "t50_h": "T50(h)",
            "mean_germination_time_h": "平均发芽时间(h)",
            "germination_speed_index": "发芽速度指数",
            "uniformity_score": "整齐度",
            "vigor_score": "活力评分",
        },
    )
    st.dataframe(comparison_view, hide_index=True, use_container_width=True)

    if not comparison.empty:
        best = comparison.iloc[0]
        cols = st.columns(4)
        cols[0].metric("最优批次", str(best["batch"]))
        cols[1].metric("活力评分", f"{float(best['vigor_score']):.1f} / 100")
        cols[2].metric("最终发芽率", f"{float(best['final_germination_rate']):.1f}%")
        cols[3].metric("T50", "未达到" if pd.isna(best["t50_h"]) else f"{float(best['t50_h']):.1f} h")

    st.subheader("发芽率趋势")
    pivot = details.pivot_table(index="time_h", columns="batch", values="germination_rate", aggfunc="last").sort_index()
    st.line_chart(pivot)

    st.subheader("时间点明细")
    details_view = pretty_frame(
        details,
        {
            "batch": "批次",
            "time_h": "时间(h)",
            "filename": "文件名",
            "total": "种子总数",
            "germinated": "已发芽",
            "non_germinated": "未发芽",
            "germination_rate": "发芽率(%)",
        },
    )
    st.dataframe(details_view, hide_index=True, use_container_width=True)

    package = build_report_package(experiment_name, comparison, details)
    st.download_button(
        "下载实验报告包（CSV + Markdown + Word）",
        data=package,
        file_name=f"{experiment_name}_report.zip",
        mime="application/zip",
    )


def render_model_results() -> None:
    st.header("模型训练与真实测试集评估")
    det = read_csv_if_exists(str(RESULTS_DIR / "detection_metrics.csv"))
    germ = read_csv_if_exists(str(RESULTS_DIR / "germination_metrics.csv"))
    t50 = read_csv_if_exists(str(RESULTS_DIR / "t50_summary.csv"))

    if det is None and germ is None and t50 is None:
        st.info("尚未接入训练评估结果。")
        return

    metric_cols = st.columns(5)
    if det is not None and not det.empty:
        row = det.iloc[0]
        metric_cols[0].metric("mAP50", f"{float(row['mAP50']):.3f}")
        metric_cols[1].metric("mAP50-95", f"{float(row['mAP50_95']):.3f}")
        metric_cols[2].metric("Precision", f"{float(row['precision']):.3f}")
        metric_cols[3].metric("Recall", f"{float(row['recall']):.3f}")
    if germ is not None and not germ.empty:
        row = germ.iloc[0]
        metric_cols[4].metric("发芽率 MAE", f"{float(row['germination_rate_MAE']) * 100:.2f}%")

    if t50 is not None and not t50.empty:
        row = t50.iloc[0]
        st.caption(
            f"T50 评估：{int(row['num_sequences'])} 组时序样本，"
            f"{int(row['valid_t50_sequences'])} 组可计算 T50，"
            f"MAE = {float(row['t50_MAE_h']):.2f} h，"
            f"RMSE = {float(row['t50_RMSE_h']):.2f} h。"
        )

    tabs = st.tabs(["检测指标", "发芽率误差", "T50", "训练曲线"])
    with tabs[0]:
        st.dataframe(det if det is not None else pd.DataFrame(), hide_index=True, use_container_width=True)
    with tabs[1]:
        st.dataframe(germ if germ is not None else pd.DataFrame(), hide_index=True, use_container_width=True)
    with tabs[2]:
        t50_full = read_csv_if_exists(str(RESULTS_DIR / "t50_metrics.csv"))
        st.dataframe(t50_full if t50_full is not None else pd.DataFrame(), hide_index=True, use_container_width=True)
    with tabs[3]:
        curve = read_csv_if_exists(str(TRAINING_RESULTS))
        if curve is None or curve.empty:
            st.info("未找到训练日志。")
        else:
            columns = [col for col in ["metrics/mAP50(B)", "metrics/mAP50-95(B)"] if col in curve.columns]
            if columns:
                st.line_chart(curve.set_index("epoch")[columns])
            st.dataframe(curve.tail(10), hide_index=True, use_container_width=True)


model_status = "已接入训练权重" if MODEL_PATH.exists() else "未发现训练权重"
detector_options = ["YOLO 训练检测器"] if MODEL_PATH.exists() else []
detector_options.append("OpenCV 规则基线")

with st.sidebar:
    st.subheader("推理设置")
    detector_choice = st.radio("检测器", detector_options)
    precision_preset = st.radio("识别策略", ["高精度", "平衡", "高召回"], index=0)
    default_conf = {"高精度": 0.65, "平衡": 0.45, "高召回": 0.25}[precision_preset]
    conf = st.slider(
        "YOLO 置信度阈值",
        0.05,
        0.90,
        default_conf,
        0.05,
        disabled=not detector_choice.startswith("YOLO"),
        help="阈值越高，误检越少，但可能漏掉不清晰目标。",
    )
    iou = st.slider(
        "YOLO NMS IoU",
        0.10,
        0.90,
        0.45,
        0.05,
        disabled=not detector_choice.startswith("YOLO"),
        help="较低 IoU 会更积极地合并重叠框。",
    )
    conservative_germination = st.checkbox(
        "保守发芽判定",
        value=precision_preset == "高精度",
        disabled=not detector_choice.startswith("YOLO"),
        help="对低置信度且形态不明显的“已发芽”结果降为未发芽，减少发芽率虚高。",
    )
    with st.expander("高级判定阈值"):
        germination_conf_threshold = st.slider("发芽类别确认阈值", 0.50, 0.95, 0.72, 0.01)
        germination_aspect_threshold = st.slider("形态拉长确认阈值", 1.10, 3.00, 1.60, 0.05)

    raw_analyzer, analyzer_name = get_analyzer(detector_choice, conf, iou)

    def analyzer(image: Image.Image) -> ImageAnalysis:
        return postprocess_analysis(
            raw_analyzer(image),
            conservative_germination=conservative_germination,
            germination_conf_threshold=germination_conf_threshold,
            germination_aspect_threshold=germination_aspect_threshold,
        )

    st.caption(f"当前推理模式：{analyzer_name}")

    st.subheader("输入模式")
    mode = st.radio(
        "分析模式",
        [
            "内置时序示例",
            "上传单张图像",
            "本机摄像头拍照",
            "摄像头定时监测",
            "上传时序图像",
            "多批次对比与报告",
            "人工校正与数据闭环",
            "训练评估结果",
        ],
    )
    interval_hours = st.number_input("时序拍摄间隔（小时）", min_value=0.5, max_value=24.0, value=6.0, step=0.5)


render_app_header(model_status, analyzer_name, mode)


if mode == "内置时序示例":
    if not SAMPLE_META.exists():
        st.warning("尚未生成内置示例。请先运行 `python scripts/make_sample_sequence.py`。")
    else:
        rows = load_sample_metadata(SAMPLE_META)
        st.header("时序培养皿图像")
        selected = st.select_slider(
            "选择时间点",
            options=list(range(len(rows))),
            format_func=lambda index: f"{rows[index]['time_h']} h",
        )
        image = Image.open(rows[selected]["path"]).convert("RGB")
        show_single_image_analysis(image, f"示例时间点：{rows[selected]['time_h']} h", analyzer)
        st.header("发芽曲线与活力评估")
        render_vigor_summary(build_summary_table(rows, analyzer=analyzer))

elif mode == "上传单张图像":
    st.header("上传单张培养皿图像")
    uploaded = st.file_uploader("上传 jpg/png/bmp/webp 图像", type=["jpg", "jpeg", "png", "bmp", "webp"])
    if uploaded is None:
        st.info("上传后系统会运行检测模型，并统计已发芽与未发芽种子数量。")
    else:
        show_single_image_analysis(Image.open(uploaded).convert("RGB"), uploaded.name, analyzer)

elif mode == "本机摄像头拍照":
    st.header("本机摄像头拍照")
    shot = st.camera_input("将培养皿放入固定视野后拍照")
    if shot is None:
        st.info("拍照后系统会直接调用当前检测器统计发芽率。")
    else:
        show_single_image_analysis(Image.open(shot).convert("RGB"), "摄像头拍照图像", analyzer)

elif mode == "摄像头定时监测":
    render_timed_camera_page(analyzer)

elif mode == "上传时序图像":
    st.header("上传一组培养皿时序图像")
    uploads = st.file_uploader(
        "上传多张图像，系统按文件名排序，并按拍摄间隔生成时间轴",
        type=["jpg", "jpeg", "png", "bmp", "webp"],
        accept_multiple_files=True,
    )
    if not uploads:
        st.info("建议文件名包含时间顺序，例如 `plate_000.png`, `plate_001.png`, `plate_002.png`。")
    else:
        rows = sequence_rows_from_uploads(uploads, interval_hours)
        summary_records = []
        tabs = st.tabs([f"{row['time_h']} h" for row in rows])
        for tab, row in zip(tabs, rows):
            with tab:
                result = show_single_image_analysis(row["image"], row["name"], analyzer)
                summary_records.append(
                    {
                        "time_h": row["time_h"],
                        "total": result.total,
                        "germinated": result.germinated,
                        "non_germinated": result.non_germinated,
                        "germination_rate": round(result.germination_rate, 2),
                    }
                )
        st.header("时序统计")
        render_vigor_summary(pd.DataFrame(summary_records))

elif mode == "多批次对比与报告":
    render_batch_comparison_page(analyzer, interval_hours)

elif mode == "人工校正与数据闭环":
    render_correction_page(analyzer)

elif mode == "训练评估结果":
    render_model_results()
